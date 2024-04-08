import json 
from flask import Flask, request, jsonify
import os
from pynput.keyboard import Key,Controller
from pynput.mouse import Button, Controller

import webbrowser
# import mediapipe as mp
import time
import numpy as np
import docx
import models as jazzModels
import pyperclip
from flask_cors import CORS



app = Flask(__name__)
CORS(app, origins='http://localhost:3000')  # Enable CORS for all routes


keyboard = Controller()
mouse = Controller()


@app.route("/open_app")
def open_app():
    text = request.args.get("text")
    appPath = f"/System/Applications/{text}.app"
    if (os.system(f"open {appPath}")) != 0: # App is not a system app
        print("Trying /Application")
        appPath = f"/Applications/{text}.app"

    return f"Launching app {appPath}"

@app.route("/close_app")
def close_app():
    appName = request.args.get("text")
    print(f"App name: {appName}")
    if str.split(appName, " ")[0].lower() == "close":
        appName = str(str.split(appName, " ")[1:])
    if str.split(appName, " ")[-1].lower() == "app":
        print("here")
        name = (str.split(appName, " ")[:-1])
        appName = ""
        for i in name:
            appName += i + " "

    appName = appName.strip()
    appName = str.replace(appName, " ", "\\ ")
    appName = str(appName)
    print(f"App name: {appName}")
    print(type(appName))
    if os.system(f"killall {appName}") != 0: # try capitalize app name 
        appName = appName.title()
        os.system(f"killall {appName}")
    return "Closing app"

@app.route("/vol_up")
def vol_up():
    keyboard.press(Key.media_volume_up)
    keyboard.release(Key.media_volume_up)
    return "volume increased"

@app.route("/vol_down")
def vol_down():
    keyboard.press(Key.media_volume_down)
    keyboard.release(Key.media_volume_down)
    return "volume decreased"


@app.route("/scroll_up")
def scroll_up():
    # Scroll two steps up
    mouse.scroll(0, -2)
    return "scroll up"

@app.route("/scroll_down")
def scroll_down():
    # Scroll two steps down
    mouse.scroll(0, 2)
    return "scroll down"


@app.route("/move_mouse", methods=['POST'])
def move_mouse():
    SCREEN_WIDTH = 2560
    SCREEN_HEIGHT = 1600
    MOUSE_SENSITIVITY = 2  # Adjust sensitivity as needed
    data = request.get_json()
    x_sum = data.get('x')
    y_sum = data.get('y')
    
    x_avg = x_sum / 21
    y_avg = y_sum / 21

    # Scale the average motion to fit within the screen dimensions
    x_movement = x_avg * MOUSE_SENSITIVITY
    y_movement = y_avg * MOUSE_SENSITIVITY

    # # Calculate new mouse position
    current_x, current_y = mouse.position
    # new_x = max(0, min(SCREEN_WIDTH - 1, current_x + x_movement))
    # new_y = max(0, min(SCREEN_HEIGHT - 1, current_y + y_movement))
    new_x =  x_movement
    new_y =  y_movement

    # Move the mouse
    mouse.position = (new_x, new_y)
    print("Current pos", mouse.position)

    print(f"x: {x_avg} y: {y_avg}")
    return {"message": "Mouse moved successfully"}




def other(output_json):
    global event_stream
    # event_stream.append(output_json["jazz_output"])
    jazzModels.speak(output_json["jazz_output"])

def type_text(text):
    pyperclip.copy(text)
    with keyboard.pressed(Key.cmd):
        keyboard.press('v')
        keyboard.release('v')
    
def goto_link(link):
    if ("http://www." in link) or ("https://www." in link):
        check = ""
        if ("http://www." in link):
            check = "http://www."
        else:
            check = "https://www."
        
        if str.replace(link, check, "").__contains__(""):
            link += ".com"
    else:
        if "." not in link:
            link += ".com"
        link = "https://www." + link


    webbrowser.open(link)
    messageHistory.append(f"Jazz: Launching {link}")
    return f"Launching {link}"

def get_word_doc_text(filename):
    doc = docx.Document(filename)
    all_paras = doc.paragraphs
    content = ""
    for para in all_paras:
       content += "\n" + para.text
    return content

def get_word_doc_summary(filename):
    doc = docx.Document(filename)
    all_paras = doc.paragraphs
    content = ""
    for para in all_paras:
       content += "\n" + jazzModels.summary(para.text)
    return content

def add_to_word_doc(filename, content):
    try:
        doc = docx.Document(filename)
        doc.add_paragraph(content)
        doc.save(filename)
        return True
    except:
        return False

def write_to_word_doc(filename, content):
    mydoc = docx.Document()
    mydoc.add_paragraph(content)
    mydoc.save(filename)

def create_new_file(filename):

    path = os.path.join("./", filename)
    os.system(f"touch {path}")
    return f"Creating {filename}"
     
def move_file(filename, location):
    filename = request.args.get('fileName')
    path = os.path.join(location, filename)
    os.system(f"mv {filename} {location}")
    return f"Moving {filename}"
 
def append_to_file():

    filename = request.args.get('fileName')
    location = request.args.get('location')
    additions = request.args.get('additions')
#    path = os.path.join(location, filename)
#    os.system(f"touch {path}")]
    
    return f"Appending to {filename}"

def getFileContents(path):
    file = open(path, "r")
    file_contents = file.read()
    file.close()
    return file_contents
    
def add_to_file(filename, changesRequested):
    global rules
    # location = "./templates/"
    location = ""
    
    path = os.path.join(location, filename)
#   read file contents
    print(f"filename: {filename}")

    fileType = str.split(filename, ".")[1]
    if fileType == "docx":
        old_contents = get_word_doc_text(path)
    else:
        old_contents = getFileContents(path)
    # code to reduce text size
    contents_arr = str.split(old_contents, "\n")

    counter = 0
    for line in contents_arr:
        line_arr = (str.split(line, " "))
        if len(line_arr) >= 10:
            # print(line_arr)
            print(f"\nLine: {line}\n")
            contents_arr[counter] = "[SUMMARY] " + jazzModels.summary(line) + " ..."
        else:
            contents_arr[counter] = line
        counter += 1

    old_contents = ""
    for line in contents_arr:
        old_contents += line + "\n"
    
    # changesRequested = jazz("Give me a complete statement telling me what additions the user wants to make to their file without the filename i.e 'add a conclusion to the file / add a paragraph about love': " + changesRequested) 

    requestString = f"""
     {changesRequested}.

    Where you encounter [SUMMARY] just know some text has been left out for brevity. 
    Complete the document for me based on the previous contents, add this after 'Addition':
    
    {filename}:
    {old_contents}

    Addition:\n
    """

    print(f"Request String: \n{requestString}")

    new_file_contents = jazzModels.jazz(requestString)

    print(f"New file contents: {new_file_contents}")

    update_file_state["filename"] = filename
    update_file_state["old_content"] = old_contents
    update_file_state["new_content"] = new_file_contents
    return [old_contents, new_file_contents]
 
def update_file(filename, changesRequested):
    global update_file_state
    global jazz_state

    location = ""
    
    path = os.path.join(location, filename)
#   read file contents
    print(f"filename: {filename}")

    fileType = str.split(filename, ".")[1]
    if fileType == "docx":
        old_contents = get_word_doc_text(path)
    else:
        old_contents = getFileContents(path)

    print(old_contents)
    requestString = f"""
    You are an expert text editor. The text you have to update will be under the "TEXT:" section. These are the instructions for what to do to the text: {changesRequested}. Do not give explanations / personal response / any boilerplate text, just output the text with the requested changes.

    TEXT:
    {old_contents}
    """
    # print(requestString)
    new_file_contents = jazzModels.jazz(requestString)
    jazz_state["update_file"]["old_file"] = old_contents
    jazz_state["update_file"]["new_file"] = new_file_contents
    # util.updateJazzStateDict(jazz_state)
  
    return [old_contents, new_file_contents]


# @app.route('/transcribe_audio', methods=['POST', 'OPTIONS'])
# def transcribe_audio():
#     transcript = jazzModels.transcribe_webm_file()  
#     print("\nThis is the transcript from FLASK:\n", transcript)
#     resp = jsonify({"transcript": transcript})
#     return resp


@app.route("/transcribe_audio")
def transcribe_audio():
    transcript = jazzModels.transcribe_webm_file()  
    print("\nThis is the transcript from FLASK:\n", transcript)
    resp = jsonify({"transcript": transcript})
    resp.headers.add('Access-Control-Allow-Origin', '*')
    print("resp:: " ,resp)
    return resp
     

@app.route("/confirm_file_update")
def write_to_file():
    filename = update_file_state["filename"]
    content = update_file_state["new_content"]
    location = ""
    fileType = str.split(filename, ".")[1]
    
    path = os.path.join(location, filename)
    print(f"\nConfirm write to {filename}")
    if confirmation():
        if not jazzModels.play_audio("write_to_file.mp3"):
            jazzModels.createSpeechFile("write_to_file.mp3", "Writing to file")
        if fileType == "docx":
            write_to_word_doc(filename, content)
        else:
            file = open(path, "w")
            file.write(content)
            file.close()
        return f"Wrote to {path}"
    else:
        return "FIle not written to "

def move_audio_to_file():
    # Move the recorded audio from downloads to /saved_responses dir
    attemps = 10
    
    move_str = f"mv /Users/seandzokesayi/Downloads/{recording_name} ./saved_responses/{recording_name}"
    if os.system(move_str) != 0:
        time.sleep(3)
        os.system(move_str)

@app.route('/processUserAudio')
def processUserAudio():
    try:
        move_audio_to_file()

        # Example usage:
        input_file = "./saved_responses/user_recorded_audio.mp3"
        output_file = "./saved_responses/updated_user_recorded_audio.mp3"

        # recording saved from browser is damaged, convet it into a new mp3
        os.system(f"ffmpeg -i {input_file} {output_file}")
        
        audio_transcript = jazzModels.transcribe_mp3_file(f"{output_file}")
        # clean up. remove audio files
        os.system(f"rm {input_file}")
        os.system(f"rm {output_file}")
    
        return audio_transcript
    except:
        return  "ERROR: 67"   

if __name__ == "__main__":
    # util.set_interval(updateMessagesOnChange, 1)
    app.run(debug=True, host="0.0.0.0", port=5001)
